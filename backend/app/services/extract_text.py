"""Extracción de texto desde archivos de requerimiento (PDF / Word / texto).

Se usa para que el consultor suba el documento funcional y el sistema lea su contenido.
"""
from __future__ import annotations
import io


class UnsupportedFile(ValueError):
    pass


def extract(filename: str, content: bytes) -> str:
    name = (filename or "").lower()

    if name.endswith((".txt", ".md", ".csv", ".log", ".abap")):
        return content.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError as e:  # pragma: no cover
            raise UnsupportedFile("Soporte PDF no instalado (pypdf).") from e
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    if name.endswith(".docx"):
        try:
            import docx  # python-docx
        except ImportError as e:  # pragma: no cover
            raise UnsupportedFile("Soporte Word no instalado (python-docx).") from e
        doc = docx.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts).strip()

    raise UnsupportedFile("Formato no soportado. Usa .pdf, .docx, .txt o .md.")
